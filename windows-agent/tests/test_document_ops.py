"""Milestone 6 — document.* executor + replace_text verifier.

Two layers of tests:
  1. The DocumentExecutor in isolation (each read action's success + failure
     paths, and replace_text: single-run formatting preservation, multiple
     occurrences, the `count` limit, text_not_found, save_as vs in-place).
  2. The full pipeline via the Dispatcher (AllowAllPolicy + document verifiers),
     proving execution AND independent re-scan verification agree end-to-end.

`.docx` fixtures are built on the fly with python-docx in `tmp_path` (pytest
builtin), so no binary fixtures are committed and the real python-docx
load/save path is exercised safely.
"""

import datetime as dt
from pathlib import Path

from docx import Document

from windows_agent.contracts import Action, ActionStatus, ExecutorResult, VerificationStatus
from windows_agent.execution import ActionRegistry, Dispatcher
from windows_agent.executors.document_ops import DocumentExecutor, register_document_executor
from windows_agent.policy import AllowAllPolicy
from windows_agent.verification import (
    DocumentReplaceTextVerifier,
    VerificationRegistry,
    register_document_verifiers,
)


def _action(type_: str, target=None, parameters=None) -> Action:
    return Action(
        action_id="a1",
        task_id="t1",
        sequence=0,
        type=type_,
        target=str(target) if target is not None else None,
        parameters=parameters or {},
        reason="test",
    )


def _make_document(path: Path) -> None:
    """Build a small known document: heading, paragraphs (one bold run), a table."""
    doc = Document()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("The quick brown fox.")
    p = doc.add_paragraph()
    run = p.add_run("Status: old value")
    run.bold = True
    doc.add_paragraph("")  # an empty paragraph (skipped by read_text join)
    doc.add_paragraph("The fox jumps over the fox.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell old"
    table.cell(0, 1).text = "other"
    doc.save(str(path))


# ── DocumentExecutor unit tests: read-only ─────────────────────────────────
async def test_read_text(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.read_text", doc_path))
    assert res.success is True
    text = res.evidence["text"]
    assert "Report Title" in text
    assert "The quick brown fox." in text
    # Empty paragraphs are not joined, so no doubled blank lines.
    assert "\n\n" not in text
    assert res.evidence["truncated"] is False
    assert res.evidence["paragraph_count"] >= 5


async def test_read_text_truncation(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.read_text", doc_path, {"max_chars": 5}))
    assert res.success is True
    assert res.evidence["truncated"] is True
    assert len(res.evidence["text"]) == 5


async def test_get_metadata(tmp_path):
    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("body")
    doc.core_properties.title = "My Title"
    doc.core_properties.author = "Ada"
    doc.core_properties.subject = "Testing"
    doc.core_properties.keywords = "alpha,beta"
    created = dt.datetime(2024, 1, 2, 3, 4, 5)
    doc.core_properties.created = created
    doc.core_properties.last_modified_by = "Grace"
    doc.save(str(doc_path))

    ex = DocumentExecutor()
    res = await ex.execute(_action("document.get_metadata", doc_path))
    assert res.success is True
    meta = res.evidence["metadata"]
    assert meta["title"] == "My Title"
    assert meta["author"] == "Ada"
    assert meta["subject"] == "Testing"
    assert meta["keywords"] == "alpha,beta"
    assert meta["last_modified_by"] == "Grace"
    # Datetime is an ISO string; compare on the stable prefix (tz suffix, if any,
    # is added by python-docx's W3CDTF serialisation and is not asserted here).
    assert isinstance(meta["created"], str)
    assert meta["created"].startswith("2024-01-02T03:04:05")
    # Unset string property is reported as null, not "".
    assert meta["content_status"] is None


async def test_find_per_paragraph_counts(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.find", doc_path, {"query": "fox"}))
    assert res.success is True
    # "The quick brown fox." (1) and "The fox jumps over the fox." (2) → 3 total.
    counts = {m["count"] for m in res.evidence["matches"]}
    assert res.evidence["total_matches"] == 3
    assert 2 in counts and 1 in counts
    assert res.evidence["truncated"] is False


async def test_find_empty_query_fails(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.find", doc_path, {"query": ""}))
    assert res.success is False
    assert res.error.code == "invalid_parameters"


# ── DocumentExecutor unit tests: replace_text ──────────────────────────────
async def test_replace_text_preserves_run_formatting(tmp_path):
    """A single-run replacement keeps that run's bold/italic formatting."""
    doc_path = tmp_path / "fmt.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Status: old value")
    run.bold = True
    run.italic = True
    doc.save(str(doc_path))

    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    )
    assert res.success is True
    assert res.evidence["replacements"] == 1
    assert res.evidence["save_as"] is False
    assert res.evidence["output_path"] == str(doc_path)

    reopened = Document(str(doc_path))
    para = reopened.paragraphs[0]
    assert para.text == "Status: new value"
    # Formatting survives on the replaced run.
    target = next(r for r in para.runs if "new" in r.text)
    assert target.bold is True
    assert target.italic is True


async def test_replace_text_multiple_occurrences_and_tables(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    # "old" appears in the bold status paragraph AND in the table cell.
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "old", "replace": "NEW"})
    )
    assert res.success is True
    assert res.evidence["replacements"] == 2
    reopened = Document(str(doc_path))
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    all_text += "\n" + reopened.tables[0].cell(0, 0).text
    assert "old" not in all_text
    assert all_text.count("NEW") == 2


async def test_replace_text_count_limit(tmp_path):
    doc_path = tmp_path / "count.docx"
    doc = Document()
    doc.add_paragraph("foo foo foo")
    doc.save(str(doc_path))
    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "foo", "replace": "bar", "count": 2})
    )
    assert res.success is True
    assert res.evidence["replacements"] == 2
    reopened = Document(str(doc_path))
    assert reopened.paragraphs[0].text == "bar bar foo"


async def test_replace_text_not_found(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "absent-string", "replace": "x"})
    )
    assert res.success is False
    assert res.error.code == "text_not_found"


async def test_replace_text_empty_find_fails(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "", "replace": "x"})
    )
    assert res.success is False
    assert res.error.code == "invalid_parameters"


async def test_replace_text_save_as_leaves_original_untouched(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("keep old text")
    doc.save(str(src))

    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", src, {"find": "old", "replace": "new", "save_as": str(out)})
    )
    assert res.success is True
    assert res.evidence["save_as"] is True
    assert res.evidence["output_path"] == str(out)
    # Original untouched; new file has the correction.
    assert Document(str(src)).paragraphs[0].text == "keep old text"
    assert Document(str(out)).paragraphs[0].text == "keep new text"


async def test_replace_text_in_place_edits_original(tmp_path):
    doc_path = tmp_path / "inplace.docx"
    doc = Document()
    doc.add_paragraph("edit old here")
    doc.save(str(doc_path))
    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    )
    assert res.success is True
    assert Document(str(doc_path)).paragraphs[0].text == "edit new here"


async def test_replace_text_save_as_existing_fails_without_overwrite(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "existing.docx"
    Document().save(str(out))  # pre-existing target
    doc = Document()
    doc.add_paragraph("has old text")
    doc.save(str(src))

    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", src, {"find": "old", "replace": "new", "save_as": str(out)})
    )
    assert res.success is False
    assert res.error.code == "output_exists"


async def test_replace_text_cross_run_fallback_collapses_formatting(tmp_path):
    """A match spanning runs collapses to the first run (documented limitation)."""
    doc_path = tmp_path / "crossrun.docx"
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hel")
    r1.bold = True
    r2 = p.add_run("lo world")  # not bold
    doc.save(str(doc_path))

    ex = DocumentExecutor()
    res = await ex.execute(
        _action("document.replace_text", doc_path, {"find": "Hello", "replace": "Goodbye"})
    )
    assert res.success is True
    assert res.evidence["replacements"] == 1
    reopened = Document(str(doc_path))
    assert reopened.paragraphs[0].text == "Goodbye world"


# ── Error paths ─────────────────────────────────────────────────────────────
async def test_read_missing_file(tmp_path):
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.read_text", tmp_path / "nope.docx"))
    assert res.success is False
    assert res.error.code == "file_not_found"


async def test_non_docx_file_fails(tmp_path):
    fake = tmp_path / "data.txt"
    fake.write_text("not a document")
    ex = DocumentExecutor()
    res = await ex.execute(_action("document.read_text", fake))
    assert res.success is False
    assert res.error.code == "not_a_document"


# ── Verifier unit tests ────────────────────────────────────────────────────
async def test_replace_verifier_passes(tmp_path):
    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("correct the old spelling")
    doc.save(str(doc_path))
    ex = DocumentExecutor()
    action = _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    res = await ex.execute(action)
    assert res.success is True
    vr = await DocumentReplaceTextVerifier().verify(action, res)
    assert vr.status == VerificationStatus.PASSED
    assert vr.observed["replace_count"] >= 1
    assert vr.observed["find_count"] == 0


async def test_replace_verifier_fails_when_replacement_absent(tmp_path):
    """Evidence claims a replacement that is not actually present → FAILED."""
    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("nothing to see")
    doc.save(str(doc_path))
    action = _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    # Fabricated evidence: says it replaced once, but "new" is nowhere on disk.
    fake = ExecutorResult(
        success=True,
        evidence={"output_path": str(doc_path), "find": "old", "replace": "new", "replacements": 1},
    )
    vr = await DocumentReplaceTextVerifier().verify(action, fake)
    assert vr.status == VerificationStatus.FAILED


async def test_replace_verifier_fails_when_reverted(tmp_path):
    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("fix the old bug")
    doc.save(str(doc_path))
    ex = DocumentExecutor()
    action = _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    res = await ex.execute(action)
    # Externally revert the document so re-observation disagrees with intent.
    reverted = Document()
    reverted.add_paragraph("fix the old bug")
    reverted.save(str(doc_path))
    vr = await DocumentReplaceTextVerifier().verify(action, res)
    assert vr.status == VerificationStatus.FAILED


# ── End-to-end via the Dispatcher (policy + execution + verification) ───────
def _pipeline():
    reg = ActionRegistry()
    register_document_executor(reg)
    vreg = VerificationRegistry()
    register_document_verifiers(vreg)
    return Dispatcher(reg, AllowAllPolicy(), verification=vreg)


async def test_pipeline_replace_text_success_and_verified(tmp_path):
    doc_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("please correct old text here")
    doc.save(str(doc_path))
    disp = _pipeline()
    result = await disp.dispatch(
        _action("document.replace_text", doc_path, {"find": "old", "replace": "new"})
    )
    assert result.status == ActionStatus.SUCCESS
    assert result.verification.status == VerificationStatus.PASSED
    assert Document(str(doc_path)).paragraphs[0].text == "please correct new text here"


async def test_pipeline_read_text_skips_verification(tmp_path):
    doc_path = tmp_path / "doc.docx"
    _make_document(doc_path)
    disp = _pipeline()
    result = await disp.dispatch(_action("document.read_text", doc_path))
    assert result.status == ActionStatus.SUCCESS
    # Read-only → no verifier registered → SKIPPED.
    assert result.verification.status == VerificationStatus.SKIPPED


async def test_register_document_executor_covers_all_types():
    reg = ActionRegistry()
    register_document_executor(reg)
    for t in (
        "document.read_text",
        "document.get_metadata",
        "document.find",
        "document.replace_text",
    ):
        assert reg.has_action(t)
