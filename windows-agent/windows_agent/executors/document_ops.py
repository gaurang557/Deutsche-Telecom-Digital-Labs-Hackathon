"""`DocumentExecutor` — the `.docx` Word-document executor (Milestone 6).

WHAT THIS IS
------------
A single executor that performs the `document.*` semantic actions against
`.docx` documents via **python-docx** (`from docx import Document`) — a
structured document API, far more reliable than scraping a viewer's GUI (see
ARCHITECTURE §3 executor-preference order). Like `FileExecutor` /
`PdfExecutor` / `SpreadsheetExecutor`, it is deliberately "dumb" about safety:
it does NOT decide permissions or ask for confirmation. That is the Policy
Engine's job (the dispatcher only runs this executor after an ALLOW decision).
See `executors/base.py`.

ACTION VOCABULARY (this milestone)
----------------------------------
Read-only (no side effects, verifier SKIPPED):
  * document.read_text    — the document's text (non-empty paragraphs joined
                            with "\n"), bounded.
  * document.get_metadata — the document's core properties (title/author/…).
  * document.find         — per-paragraph case-sensitive substring hit counts.

Modifying (verified by re-observation in
`verification/document_verifiers.py`):
  * document.replace_text — replace occurrences of `find` with `replace`,
                            **preserving formatting**, across body paragraphs,
                            table cells, and section headers/footers. The
                            headline M6 capability: correcting a document
                            without disturbing its formatting.

PARAMETER CONVENTIONS
---------------------
`action.target` is the PRIMARY path (the document `.docx`). Everything else
lives in `action.parameters`, e.g. {"max_chars": 5000}, {"query": "foo"},
{"find": "old", "replace": "new", "count": 2, "save_as": "out.docx"}. Only
`.docx` is supported (the legacy binary `.doc` format is not).

FORMATTING PRESERVATION — approach + its documented limitation
--------------------------------------------------------------
`.docx` stores a paragraph as a sequence of *runs*, and each run carries its own
formatting (bold/italic/font/…). Word freely splits a logical piece of text
across several runs, so a search string may live entirely inside one run or
straddle a run boundary.

We handle the two cases pragmatically:
  * **Match within a single run (the clean case):** we replace the substring
    *inside that run's text only*, so the run — and therefore all of its
    formatting — is preserved exactly. This is the common case and the one we
    assert on in tests.
  * **Match spanning multiple runs (the fallback):** there is no single run
    whose formatting is "the" formatting of the match, so we rebuild the
    paragraph text across its runs, apply the replacement, write the whole
    result into the paragraph's FIRST run, and clear the remaining runs.
    **LIMITATION:** a cross-run replacement therefore collapses the affected
    text to the first run's formatting. This is an accepted, documented
    trade-off (a fully faithful cross-run edit would require splitting/merging
    runs at XML level, which is out of scope for M6).

WHY ASYNC + to_thread
---------------------
The executor contract is async, but python-docx's load/save calls are blocking
(they parse/serialise a zip of XML on disk). Each operation is dispatched to a
worker thread via `asyncio.to_thread` instead of blocking the event loop —
exactly as file_ops/pdf_ops/spreadsheet_ops do for their blocking I/O.

SAFETY / BOUNDING
-----------------
  * `document.read_text` is capped at `_DEFAULT_TEXT_CHAR_CAP` characters
    (override with `max_chars`) and `document.find` at `_DEFAULT_SEARCH_RESULTS`
    matching-paragraph entries, so a huge document can never be slurped into
    memory / evidence; `truncated=true` marks a clipped result. The dispatcher
    additionally bounds evidence.
  * `document.replace_text` with an empty `find` is rejected; a `find` that is
    absent everywhere fails with `text_not_found` (0 replacements is reported as
    an ERROR, so the planner learns the correction did not apply).
  * `save_as` writes the result to a NEW path (the original is untouched);
    without it the executor edits the document IN PLACE (overwriting the
    original). A `save_as` that would clobber an existing DIFFERENT file fails
    closed with `output_exists` unless `overwrite=true`.
  * Expected errors are returned as `ExecutorResult(success=False, error=...)` —
    the executor never raises for ordinary failures. (Unexpected exceptions are
    still contained by the dispatcher.)

RISK (documentation only — the executor NEVER sets risk)
--------------------------------------------------------
  * reads (`read_text`/`get_metadata`/`find`) → `RiskLevel.NONE`.
  * `document.replace_text` editing IN PLACE overwrites the original → `HIGH`.
  * `document.replace_text` with `save_as` writes a new file → `MEDIUM`.
See docs/ACTION_REFERENCE.md. Risk is assigned by the deterministic policy
(M12); it is documented here but never decided by this executor.
"""

from __future__ import annotations

import asyncio
import datetime as _dt
from pathlib import Path
from typing import Any, Iterator

from docx import Document

from ..contracts import Action, ActionError, ErrorCode, ExecutorResult
from .base import BaseExecutor

# Domain-specific error codes. `ActionError.code` is a free string (see
# contracts/error.py); we use stable document-specific codes here and fall back
# to the shared ErrorCode values where they fit.
ERR_FILE_NOT_FOUND = "file_not_found"
ERR_NOT_A_DOCUMENT = "not_a_document"
ERR_TEXT_NOT_FOUND = "text_not_found"
ERR_OUTPUT_EXISTS = "output_exists"
ERR_INVALID_PARAMS = "invalid_parameters"

#: Default cap on extracted text length so a huge document is never read into
#: evidence wholesale. Overridable per call via `max_chars`.
_DEFAULT_TEXT_CHAR_CAP = 20_000
#: Default/limit on the number of per-paragraph match entries `document.find`
#: returns, so a query hitting every paragraph cannot bloat evidence.
_DEFAULT_SEARCH_RESULTS = 100
#: Sentinel used when `count` is not supplied (replace ALL occurrences). A
#: document will never contain this many occurrences.
_UNLIMITED = 1_000_000_000

#: Every action type this executor handles and its deterministic verification
#: requirement.
DOCUMENT_ACTION_REQUIREMENTS: dict[str, bool] = {
    "document.read_text": False,
    "document.get_metadata": False,
    "document.find": False,
    "document.replace_text": True,
}
DOCUMENT_ACTION_TYPES: tuple[str, ...] = tuple(DOCUMENT_ACTION_REQUIREMENTS)

#: core_properties attributes we surface. Datetimes become ISO strings; empty
#: strings become null; ints (e.g. `revision`) pass through.
_METADATA_FIELDS = (
    "title",
    "author",
    "subject",
    "keywords",
    "created",
    "modified",
    "last_modified_by",
    "category",
    "comments",
    "content_status",
    "identifier",
    "language",
    "revision",
    "version",
    "last_printed",
)


def _err(code: str, message: str, *, retryable: bool = False) -> ExecutorResult:
    return ExecutorResult(
        success=False,
        error=ActionError(code=code, message=message, retryable=retryable),
    )


def _normalize_meta(value: Any) -> Any:
    """Coerce a core-property value to a JSON-serialisable primitive.

    Datetimes → ISO-8601 strings, empty strings → None, everything else (ints
    like `revision`, non-empty strings) passes through unchanged.
    """
    if value is None:
        return None
    if isinstance(value, (_dt.datetime, _dt.date, _dt.time)):
        return value.isoformat()
    if isinstance(value, str):
        return value or None
    return value


def _iter_container_paragraphs(container: Any) -> Iterator[Any]:
    """Yield every paragraph in a container (its own + those nested in tables).

    A "container" is anything exposing `.paragraphs` and `.tables`: the
    Document body, a header, a footer, or a table cell (cells can nest tables,
    so we recurse).
    """
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def _all_paragraphs(doc: Any) -> Iterator[Any]:
    """Yield every paragraph we edit/scan: body + tables + section headers/footers.

    Headers/footers linked to the previous section carry no own paragraphs, so
    shared headers are not double-processed.
    """
    yield from _iter_container_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_container_paragraphs(section.header)
        yield from _iter_container_paragraphs(section.footer)


def _replace_in_paragraph(paragraph: Any, find: str, replace: str, remaining: int) -> int:
    """Replace up to `remaining` occurrences of `find` in one paragraph.

    Returns the number of replacements performed. See the module docstring for
    the run-level (formatting-preserving) vs cross-run (fallback) strategy.
    """
    if remaining <= 0 or not paragraph.runs:
        return 0

    original_full = "".join(run.text for run in paragraph.runs)
    full_count = original_full.count(find)
    if full_count == 0:
        return 0
    # Occurrences wholly contained in a single run (non-overlapping counting,
    # so this is always <= full_count).
    within_run_count = sum(run.text.count(find) for run in paragraph.runs)

    # Clean case — every match lives inside a single run, so we can replace in
    # place and each run keeps ALL of its formatting. `str.replace` operates on
    # the original run text and never re-scans its own output, so this stays
    # correct even when `replace` contains `find` (e.g. "old" → "older").
    if within_run_count == full_count:
        done = 0
        for run in paragraph.runs:
            if remaining - done <= 0:
                break
            occurrences = run.text.count(find)
            if occurrences:
                allowed = min(occurrences, remaining - done)
                run.text = run.text.replace(find, replace, allowed)
                done += allowed
        return done

    # Fallback — at least one match straddles a run boundary. Replace against the
    # ORIGINAL joined text (so replacement output is never re-matched) and write
    # the whole result into the first run, clearing the rest. This collapses the
    # paragraph to the first run's formatting — the documented M6 limitation, and
    # only for paragraphs that actually contain a cross-run match.
    allowed = min(full_count, remaining)
    paragraph.runs[0].text = original_full.replace(find, replace, allowed)
    for run in paragraph.runs[1:]:
        run.text = ""
    return allowed


class DocumentExecutor(BaseExecutor):
    """Executes the `document.*` action vocabulary via python-docx (.docx)."""

    name = "document"

    async def execute(self, action: Action) -> ExecutorResult:
        # Route on the semantic type. The registry only sends us document.* types
        # it was told to, but we still guard against an unexpected one.
        handler = {
            "document.read_text": self._read_text,
            "document.get_metadata": self._get_metadata,
            "document.find": self._find,
            "document.replace_text": self._replace_text,
        }.get(action.type)
        if handler is None:
            return _err(
                ErrorCode.NOT_IMPLEMENTED.value,
                f"DocumentExecutor does not handle action type {action.type!r}",
            )
        # Blocking python-docx load/save runs off the event loop.
        return await asyncio.to_thread(handler, action)

    # ── helpers ────────────────────────────────────────────────────────────
    @staticmethod
    def _require_target(action: Action) -> Path | None:
        return Path(action.target) if action.target else None

    @staticmethod
    def _param(action: Action, key: str, default: Any = None) -> Any:
        return action.parameters.get(key, default)

    def _open_document(self, path: Path | None) -> tuple[Any, ExecutorResult | None]:
        """Open + guard an existing `.docx`, returning (doc, None) or
        (None, error_result).

        Fails closed on a missing path, a non-file target, a non-`.docx`
        extension, and anything python-docx cannot parse as a document.
        """
        if path is None:
            return None, _err(ERR_INVALID_PARAMS, "document action requires a target path")
        if not path.exists():
            return None, _err(ERR_FILE_NOT_FOUND, f"File not found: {path}")
        if not path.is_file():
            return None, _err(ERR_NOT_A_DOCUMENT, f"Not a file: {path}")
        if path.suffix.lower() != ".docx":
            return None, _err(ERR_NOT_A_DOCUMENT, f"Not a .docx document: {path}")
        try:
            doc = Document(str(path))
        except Exception as exc:  # python-docx raises various zip/XML/format errors
            return None, _err(ERR_NOT_A_DOCUMENT, f"Cannot open as .docx: {path}: {exc}")
        return doc, None

    # ── read-only operations ────────────────────────────────────────────────
    def _read_text(self, action: Action) -> ExecutorResult:
        path = self._require_target(action)
        doc, err = self._open_document(path)
        if err is not None:
            return err

        max_chars = self._param(action, "max_chars", _DEFAULT_TEXT_CHAR_CAP)
        try:
            max_chars = int(max_chars)
        except (TypeError, ValueError):
            return _err(ERR_INVALID_PARAMS, f"max_chars must be an integer, got {max_chars!r}")
        if max_chars <= 0:
            max_chars = _DEFAULT_TEXT_CHAR_CAP

        paragraphs = doc.paragraphs
        # Join only the non-empty body paragraphs, capping the total length.
        parts: list[str] = []
        total = 0
        truncated = False
        for para in paragraphs:
            text = para.text
            if not text:
                continue
            piece = ("\n" if parts else "") + text
            if total + len(piece) > max_chars:
                parts.append(piece[: max_chars - total])
                truncated = True
                break
            parts.append(piece)
            total += len(piece)

        return ExecutorResult(
            success=True,
            evidence={
                "path": str(path),
                "text": "".join(parts),
                "paragraph_count": len(paragraphs),
                "truncated": truncated,
            },
        )

    def _get_metadata(self, action: Action) -> ExecutorResult:
        path = self._require_target(action)
        doc, err = self._open_document(path)
        if err is not None:
            return err
        props = doc.core_properties
        metadata = {field: _normalize_meta(getattr(props, field, None)) for field in _METADATA_FIELDS}
        return ExecutorResult(
            success=True,
            evidence={"path": str(path), "metadata": metadata},
        )

    def _find(self, action: Action) -> ExecutorResult:
        path = self._require_target(action)
        query = self._param(action, "query")
        if not isinstance(query, str) or query == "":
            # Open validity is irrelevant if the query is unusable; fail fast,
            # but still surface a missing/broken file first for a clearer error.
            doc, err = self._open_document(path)
            if err is not None:
                return err
            return _err(ERR_INVALID_PARAMS, "document.find requires a non-empty parameters.query")

        doc, err = self._open_document(path)
        if err is not None:
            return err

        max_results = self._param(action, "max_results", _DEFAULT_SEARCH_RESULTS)
        try:
            max_results = int(max_results)
        except (TypeError, ValueError):
            max_results = _DEFAULT_SEARCH_RESULTS
        if max_results <= 0:
            max_results = _DEFAULT_SEARCH_RESULTS

        matches: list[dict[str, int]] = []
        total_matches = 0
        truncated = False
        for index, para in enumerate(doc.paragraphs):
            count = para.text.count(query)
            if not count:
                continue
            if len(matches) >= max_results:
                # More paragraphs match than we will report — mark bounded.
                truncated = True
                break
            matches.append({"paragraph_index": index, "count": count})
            total_matches += count

        return ExecutorResult(
            success=True,
            evidence={
                "path": str(path),
                "matches": matches,
                "total_matches": total_matches,
                "truncated": truncated,
            },
        )

    # ── modifying operation ─────────────────────────────────────────────────
    def _replace_text(self, action: Action) -> ExecutorResult:
        path = self._require_target(action)
        find = self._param(action, "find")
        replace = self._param(action, "replace", "")

        if not isinstance(find, str) or find == "":
            return _err(ERR_INVALID_PARAMS, "document.replace_text requires a non-empty parameters.find")
        if not isinstance(replace, str):
            return _err(ERR_INVALID_PARAMS, "document.replace_text parameters.replace must be a string")

        # Optional count limit: when supplied it must be a positive int.
        count_raw = self._param(action, "count")
        if count_raw is None:
            remaining = _UNLIMITED
        else:
            if isinstance(count_raw, bool) or not isinstance(count_raw, int) or count_raw <= 0:
                return _err(ERR_INVALID_PARAMS, f"count must be a positive integer, got {count_raw!r}")
            remaining = count_raw

        # Resolve the output path: save_as → a NEW file; otherwise edit in place.
        save_as = self._param(action, "save_as")
        overwrite = bool(self._param(action, "overwrite", False))
        if save_as is not None:
            if not isinstance(save_as, str) or not save_as.strip():
                return _err(ERR_INVALID_PARAMS, "save_as must be a non-empty path string")
            output_path = Path(save_as)
            if output_path.suffix.lower() != ".docx":
                return _err(ERR_INVALID_PARAMS, f"save_as must be a .docx path: {output_path}")
            if not output_path.parent.exists():
                return _err(ERR_FILE_NOT_FOUND, f"Directory does not exist: {output_path.parent}")
        else:
            output_path = path  # in-place edit (overwrites the original)

        doc, err = self._open_document(path)
        if err is not None:
            return err

        # Refuse to clobber a DIFFERENT existing file via save_as unless allowed
        # (mirrors file.write_text / spreadsheet.write_cell's fail-closed guard).
        if (
            save_as is not None
            and output_path.exists()
            and output_path.resolve() != path.resolve()
            and not overwrite
        ):
            return _err(
                ERR_OUTPUT_EXISTS,
                f"save_as target already exists (set overwrite=true to replace): {output_path}",
            )

        total = 0
        for para in _all_paragraphs(doc):
            if remaining - total <= 0:
                break
            total += _replace_in_paragraph(para, find, replace, remaining - total)

        # 0 replacements is an ERROR (text_not_found), so the planner learns the
        # correction did not apply and nothing is written.
        if total == 0:
            return _err(ERR_TEXT_NOT_FOUND, f"Text not found in document: {find!r}")

        doc.save(str(output_path))
        return ExecutorResult(
            success=True,
            evidence={
                "path": str(path),
                "output_path": str(output_path),
                "find": find,
                "replace": replace,
                "replacements": total,
                "save_as": save_as is not None,
            },
            side_effects=[{"type": "document.text_replaced", "target": str(output_path)}],
        )


def register_document_executor(
    registry, executor: DocumentExecutor | None = None, *, override: bool = False
) -> DocumentExecutor:
    """Register a single `DocumentExecutor` for every `document.*` action type.

    Returns the executor instance so callers can reuse it. One instance handles
    all document operations (it is stateless), matching the "one executor, many
    action types" pattern noted in `executors/base.py`.
    """
    executor = executor or DocumentExecutor()
    for action_type, requires_verification in DOCUMENT_ACTION_REQUIREMENTS.items():
        registry.register_action(
            action_type,
            executor,
            requires_verification=requires_verification,
            override=override,
        )
    return executor
