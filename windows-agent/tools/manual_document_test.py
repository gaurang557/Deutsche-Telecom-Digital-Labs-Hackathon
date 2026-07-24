r"""Interactive manual tester for the `document.*` functionality (Milestone 6).

This is a HUMAN-DRIVEN harness (not an automated test). It wires up the real
execution pipeline exactly as the agent will use it —

    ActionRegistry (+ DocumentExecutor)
      → Dispatcher (+ AllowAllPolicy, VerificationRegistry (+ document
        verifiers), InMemoryAuditSink)

— then lets you type document commands and shows, for each one:
  * the resulting ActionStatus,
  * the bounded evidence the executor returned (text, metadata, matches,
    replacement counts),
  * the verification result (SKIPPED for reads; an INDEPENDENT re-scan PASS/FAIL
    for `replace`/`replaceas`), and
  * the ordered audit trail emitted around the action.

Relative paths you type are resolved inside a throwaway `sandbox/` directory;
absolute paths are used as-is. Only `.docx` documents are supported.

To get started with no document of your own, run `sample` to generate a small
`sample.docx` (a heading, a few paragraphs including bold text, a table, and
core properties) inside the sandbox, then exercise the commands against it.

RUN (from the windows-agent/ folder, using the project venv):

    & "..\.venv\Scripts\python.exe" tools\manual_document_test.py
    # or point at a custom sandbox:
    & "..\.venv\Scripts\python.exe" tools\manual_document_test.py --sandbox C:\tmp\wa

Type `help` at the prompt for the command list, `quit` to exit.

NOTE ON POLICY: this harness uses AllowAllPolicy, so every action is authorized
(the real deterministic policy arrives in M12). The point here is to exercise
DocumentExecutor behaviour + the re-scan replace verifier, which is what M6
delivers. `replace` edits the file IN PLACE (HIGH risk); `replaceas` writes a
NEW file and leaves the original untouched (MEDIUM risk).
"""

from __future__ import annotations

import argparse
import asyncio
import shlex
import sys
import uuid
from pathlib import Path

# Make the module importable when run as a plain script (sys.path[0] is tools/).
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # also used to synthesise the sample document  # noqa: E402

from windows_agent.audit import InMemoryAuditSink  # noqa: E402
from windows_agent.contracts import Action  # noqa: E402
from windows_agent.execution import ActionRegistry, Dispatcher  # noqa: E402
from windows_agent.executors import register_document_executor  # noqa: E402
from windows_agent.policy import AllowAllPolicy  # noqa: E402
from windows_agent.verification import VerificationRegistry, register_document_verifiers  # noqa: E402


HELP = """
Commands (paths are relative to the sandbox unless absolute; .docx only):
  sample [name]                    generate a sample document (default: sample.docx)
  ls                               list documents in the sandbox
  read <path> [max_chars]          document.read_text    — the document's text
  meta <path>                      document.get_metadata — core properties
  find <path> <query>              document.find         — per-paragraph hit counts
  replace <path> <find> <replace> [count]
                                   document.replace_text — edit IN PLACE (HIGH)
  replaceas <path> <find> <replace> <out> [count]
                                   document.replace_text — write NEW file (MEDIUM)
  help                             show this help
  quit / exit                      leave

Try this quick scenario:
  sample
  read sample.docx
  find sample.docx old
  replace sample.docx old new
  read sample.docx
"""

# Paragraphs baked into the generated sample document (some formatted).
_SAMPLE_HEADING = "Quarterly Report"


def _make_sample(path: Path) -> tuple[int, int]:
    """Create a small document with known content. Returns (paragraph_count, bytes)."""
    doc = Document()
    doc.add_heading(_SAMPLE_HEADING, level=1)
    doc.add_paragraph("This report contains an old figure that needs correcting.")
    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run("old value").bold = True
    doc.add_paragraph("The old total will be updated in the old table below.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "old label"
    table.cell(0, 1).text = "123"
    doc.core_properties.title = "Quarterly Report"
    doc.core_properties.author = "manual_document_test"
    doc.core_properties.subject = "M6 sample"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    count = len(Document(str(path)).paragraphs)
    return count, path.stat().st_size


# Maps a typed command to (action_type, (target, params)).
def _build(cmd: str, args: list[str], resolve):
    if cmd == "read":
        params = {"max_chars": int(args[1])} if len(args) > 1 else {}
        return "document.read_text", (str(resolve(args[0])), params)
    if cmd == "meta":
        return "document.get_metadata", (str(resolve(args[0])), {})
    if cmd == "find":
        if len(args) < 2:
            raise ValueError("find needs <path> <query>")
        return "document.find", (str(resolve(args[0])), {"query": args[1]})
    if cmd == "replace":
        if len(args) < 3:
            raise ValueError("replace needs <path> <find> <replace> [count]")
        params = {"find": args[1], "replace": args[2]}
        if len(args) > 3:
            params["count"] = int(args[3])
        return "document.replace_text", (str(resolve(args[0])), params)
    if cmd == "replaceas":
        if len(args) < 4:
            raise ValueError("replaceas needs <path> <find> <replace> <out> [count]")
        params = {"find": args[1], "replace": args[2], "save_as": str(resolve(args[3]))}
        if len(args) > 4:
            params["count"] = int(args[4])
        return "document.replace_text", (str(resolve(args[0])), params)
    raise ValueError(f"Unknown command: {cmd!r} (type 'help')")


def _print_result(result, sink: InMemoryAuditSink) -> None:
    print(f"\n  status      : {result.status.value}")
    if result.evidence:
        print(f"  evidence    : {result.evidence}")
    if result.verification is not None:
        v = result.verification
        print(f"  verification: {v.status.value} ({v.method})")
        if v.message:
            print(f"                {v.message}")
        if v.expected is not None or v.observed is not None:
            print(f"                expected={v.expected!r} observed={v.observed!r}")
    if result.error is not None:
        print(f"  error       : [{result.error.code}] {result.error.message}")
    print("  audit trail :")
    for event in sink.events:
        tag = f" ({event.outcome})" if event.outcome else ""
        print(f"      - {event.event_type.value}{tag}: {event.summary}")
    print()


async def _run(sandbox: Path) -> None:
    registry = ActionRegistry()
    register_document_executor(registry)
    verification = VerificationRegistry()
    register_document_verifiers(verification)
    audit = InMemoryAuditSink()
    dispatcher = Dispatcher(registry, AllowAllPolicy(), verification=verification, audit=audit)

    task_id = uuid.uuid4().hex[:8]
    seq = 0
    resolve = lambda p: (Path(p) if Path(p).is_absolute() else sandbox / p)

    print(f"Sandbox: {sandbox}")
    print("Manual document.* tester. Type 'help' for commands, 'sample' to make a test document, 'quit' to exit.")

    while True:
        try:
            line = input("docx> ").strip()
        except (EOFError, KeyboardInterrupt):
            print()
            break
        if not line:
            continue
        if line in ("quit", "exit"):
            break
        if line == "help":
            print(HELP)
            continue

        try:
            parts = shlex.split(line)
        except ValueError as exc:  # e.g. unbalanced quotes
            print(f"  parse error: {exc}")
            continue
        cmd, args = parts[0], parts[1:]

        # Local (non-pipeline) conveniences.
        if cmd == "sample":
            name = args[0] if args else "sample.docx"
            target = resolve(name)
            count, size = _make_sample(target)
            print(f"  created {target} (paragraphs={count}, {size} bytes)")
            continue
        if cmd == "ls":
            docs = sorted(p.name for p in sandbox.glob("*.docx"))
            print(f"  {sandbox}:")
            print("    " + ("  ".join(docs) if docs else "(no documents — run 'sample')"))
            continue

        try:
            action_type, (target, params) = _build(cmd, args, resolve)
        except (ValueError, IndexError) as exc:
            print(f"  {exc}")
            continue

        action = Action(
            action_id=uuid.uuid4().hex[:8],
            task_id=task_id,
            sequence=seq,
            type=action_type,
            target=target,
            parameters=params,
            reason=f"manual: {line}",
        )
        seq += 1

        audit.events.clear()  # show only this action's events
        result = await dispatcher.dispatch(action)
        _print_result(result, audit)


def main() -> None:
    parser = argparse.ArgumentParser(description="Interactive manual tester for document.* actions")
    parser.add_argument(
        "--sandbox",
        default=str(Path.cwd() / "sandbox"),
        help="Directory to operate in (created if missing). Default: ./sandbox",
    )
    ns = parser.parse_args()
    sandbox = Path(ns.sandbox).resolve()
    sandbox.mkdir(parents=True, exist_ok=True)
    asyncio.run(_run(sandbox))


if __name__ == "__main__":
    main()
